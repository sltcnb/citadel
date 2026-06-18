"""Standalone document rendering for a single markdown body (e.g. the AI LLM
report). Turns one markdown string into HTML / DOCX so the analyst-facing
narrative can be exported as the final deliverable in any format.

PDF is produced by the API layer (WeasyPrint on the HTML from here) — kept out
of Scribe so the optional WeasyPrint system dependency stays at the edge.
"""

from __future__ import annotations

import io

from .render import _e, _HTML_CSS, _md_to_html

_DOC_CSS = _HTML_CSS + """
.docmeta{color:#64748B;font-size:.85em;margin:.2em 0 1.4em;border-bottom:1px solid #E5E7EB;padding-bottom:.8em;}
"""


def render_html_document(title: str, markdown: str, meta_lines: list[str] | None = None) -> str:
    """Full, self-contained HTML page for one markdown body."""
    body = [f"<h1>{_e(title)}</h1>"]
    if meta_lines:
        body.append('<p class="docmeta">' + " · ".join(_e(m) for m in meta_lines if m) + "</p>")
    # Demote a leading H1 inside the content (the page already has the title)
    # so we don't render two competing <h1>s.
    content = (markdown or "").strip()
    content = "\n".join(
        ("## " + ln[2:]) if ln.startswith("# ") else ln for ln in content.splitlines()
    )
    body.append(_md_to_html(content))
    page = (
        "<!doctype html><html><head><meta charset=utf-8/>"
        f"<title>{_e(title)}</title><style>{_DOC_CSS}</style>"
        f"</head><body>{''.join(body)}</body></html>"
    )
    # ASCII-safe (numeric entities for non-ASCII) — matches render_html, avoids
    # mojibake regardless of how the document is served/embedded.
    return page.encode("ascii", "xmlcharrefreplace").decode("ascii")


def render_markdown_document(title: str, markdown: str, meta_lines: list[str] | None = None) -> str:
    """Markdown body with a guaranteed title + optional metadata line."""
    content = (markdown or "").strip()
    out = []
    # Only add a title if the content doesn't already open with one.
    if not content.startswith("# "):
        out.append(f"# {title}")
        out.append("")
    if meta_lines:
        out.append("_" + " · ".join(m for m in meta_lines if m) + "_")
        out.append("")
    out.append(content)
    return "\n".join(out)


def render_docx_document(title: str, markdown: str, meta_lines: list[str] | None = None) -> bytes:
    """Render one markdown body to a .docx. Raises ImportError without python-docx."""
    from docx import Document  # type: ignore
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=0)
    if meta_lines:
        p = doc.add_paragraph(" · ".join(m for m in meta_lines if m))
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(9)
    _markdown_to_docx_body(doc, markdown or "")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _markdown_to_docx_body(doc, markdown: str) -> None:
    """Map a markdown subset (headings, bullets, ordered lists, blockquotes,
    pipe-table rows, paragraphs) onto python-docx. Inline marks are stripped to
    plain text — good enough for an editable Office handoff."""
    import re

    lines = markdown.splitlines()
    i = 0
    table_buf: list[list[str]] = []

    def _flush_table():
        nonlocal table_buf
        if not table_buf:
            return
        cols = max(len(r) for r in table_buf)
        t = doc.add_table(rows=0, cols=cols)
        try:
            t.style = "Light Grid Accent 1"
        except Exception:
            pass
        for row in table_buf:
            cells = t.add_row().cells
            for c in range(cols):
                cells[c].text = _strip_inline(row[c]) if c < len(row) else ""
        table_buf = []

    def _strip_inline(s: str) -> str:
        s = re.sub(r"`([^`]+)`", r"\1", s)
        s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
        s = re.sub(r"(?<!\w)_([^_]+)_(?!\w)", r"\1", s)
        s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)
        return s.strip()

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        # Pipe table: header row + |---| separator
        if (
            stripped.startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1])
            and "-" in lines[i + 1]
        ):
            def _cells(row: str) -> list[str]:
                return [c.strip() for c in row.strip().strip("|").split("|")]
            table_buf.append(_cells(line))
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_buf.append(_cells(lines[i]))
                i += 1
            _flush_table()
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(_strip_inline(stripped[5:]), level=4)
        elif stripped.startswith("### "):
            doc.add_heading(_strip_inline(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(_strip_inline(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(_strip_inline(stripped[2:]), level=1)
        elif stripped == "---":
            pass
        elif stripped.startswith("> "):
            p = doc.add_paragraph(_strip_inline(stripped[2:]))
            p.style = "Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else p.style
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(_strip_inline(re.sub(r"^\d+\.\s", "", stripped)), style="List Number")
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(_strip_inline(stripped[2:]), style="List Bullet")
        else:
            doc.add_paragraph(_strip_inline(stripped))
        i += 1
