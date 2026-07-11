"""DOCX rendering for Scribe reports.

Kept separate from render.py because it needs the optional `python-docx`
dependency. `render_docx` raises ImportError if it isn't installed — callers
should catch and surface a clear "DOCX export unavailable" message.

Builds the Word document from the same report `data` dict the markdown/HTML
renderers use (NOT by converting HTML), so the structure is native and clean.
"""

from __future__ import annotations

import io

from .labels import Labels
from .render import (
    TEMPLATE_DEFAULTS,
    _ev_fields,
    _ts,
)


def render_docx(data: dict, tpl: dict | None = None, language: str | None = None) -> bytes:
    """Render the report as a .docx and return the raw bytes.

    Raises ImportError when python-docx is missing.
    """
    from docx import Document  # type: ignore
    from docx.shared import Pt

    tpl = tpl or dict(TEMPLATE_DEFAULTS)
    L = Labels(language)
    sections = tpl.get("sections") or {}
    case = data.get("case") or {}
    name = case.get("name", case.get("case_id", "Case"))

    title_prefix = tpl.get("title_prefix")
    if not title_prefix or title_prefix == TEMPLATE_DEFAULTS["title_prefix"]:
        title_prefix = L("title_prefix")

    doc = Document()
    doc.add_heading(f"{title_prefix} — {name}", level=0)

    meta_bits = [f"{L('generated')} {_now_label()}"]
    if case.get("company"):
        meta_bits.append(f"{L('company')}: {case['company']}")
    doc.add_paragraph(" · ".join(meta_bits)).italic = True

    def h2(text):
        doc.add_heading(text, level=1)

    def h3(text):
        doc.add_heading(text, level=2)

    def bullet(text):
        doc.add_paragraph(str(text), style="List Bullet")

    # ── Manifest
    manifest = data.get("manifest") or {}
    if sections.get("manifest", True) and manifest:
        h2(L("manifest_title"))
        doc.add_paragraph(L("manifest_blurb"))
        if manifest.get("total_events"):
            bullet(L("m_events", n=f"{int(manifest['total_events']):,}"))
        bullet(L("m_flagged", n=manifest.get("flagged_count", 0)))
        if manifest.get("pinned_count"):
            bullet(L("m_pinned", n=manifest["pinned_count"]))
        bullet(L("m_modules", hit=manifest.get("module_hit_run_count", 0), total=manifest.get("module_run_count", 0)))
        if manifest.get("saved_search_count"):
            bullet(L("m_saved", n=manifest["saved_search_count"]))
        if manifest.get("killchain_count"):
            bullet(L("m_killchains", n=manifest["killchain_count"]))
        bullet(L("m_ai_yes", model=manifest.get("ai_model") or "?") if manifest.get("has_ai") else L("m_ai_no"))

    # ── Executive summary
    agg = data.get("aggregates") or {}
    flagged = data.get("flagged") or []
    pinned = data.get("pinned") or []
    wl = data.get("watchlist") or {}
    if sections.get("exec_summary", True):
        h2(L("exec_summary"))
        if agg.get("total_events"):
            bullet(L("total_events", n=f"{int(agg['total_events']):,}"))
        bullet(L("flagged_review", n=len(flagged)))
        if agg.get("cti"):
            bullet(L("cti_matched", n=len(agg["cti"])))
        if wl.get("hits"):
            bullet(L("wl_hits", n=len(wl["hits"])))

    # ── Module analysis
    modules = data.get("modules") or []
    if sections.get("modules", True) and modules:
        h2(L("module_analysis"))
        for r in modules[:40]:
            lv = r.get("hits_by_level") or {}
            sev = ", ".join(f"{int(n)} {k}" for k, n in lv.items() if n)
            line = f"{r.get('module_id', '?')} ({r.get('status', '')}) — {r.get('total_hits', 0)} {L('u_hits')}"
            if sev:
                line += f" [{sev}]"
            bullet(line)

    # ── Saved searches
    saved = data.get("saved_searches") or []
    if sections.get("saved_searches", True) and saved:
        h2(L("saved_title"))
        doc.add_paragraph(L("saved_blurb"))
        for s in saved:
            bullet(f"{s.get('name', '?')} — {L('saved_matches', n=int(s.get('count', 0)))}: {s.get('query', '')}")

    # ── Correlated kill chains
    killchains = data.get("killchains") or []
    if sections.get("killchains", True) and killchains:
        h2(L("killchains_title"))
        doc.add_paragraph(L("killchains_blurb"))
        for kc in killchains:
            a = kc.get("anchor") or {}
            h3(f"{L('kc_anchor')}: {a.get('summary') or a.get('fo_id', '?')}")
            if kc.get("tactics_covered"):
                doc.add_paragraph(f"{L('kc_tactics')}: {', '.join(kc['tactics_covered'])}")
            for st in kc.get("steps", [])[:30]:
                tac = st.get("tactic") or st.get("phase") or ""
                bullet(f"{_ts(st.get('ts', ''))} — {tac} {st.get('technique', '')}: {(st.get('summary') or '')[:200]}")

    # ── AI narrative (already prose) — drop in as paragraphs
    ai_report = data.get("ai_report")
    if sections.get("ai_report", True) and ai_report and (ai_report.get("content") or "").strip():
        h2(L("ai_investigation"))
        for ln in ai_report["content"].splitlines():
            ln = ln.rstrip()
            if not ln:
                continue
            if ln.startswith("### "):
                h3(ln[4:])
            elif ln.startswith("## "):
                h3(ln[3:])
            elif ln.startswith("# "):
                h3(ln[2:])
            elif ln.lstrip().startswith(("- ", "* ")):
                bullet(ln.lstrip()[2:])
            else:
                doc.add_paragraph(ln)

    # ── Pinned / flagged events
    if sections.get("pinned", True) and pinned:
        h2(L("key_evidence"))
        for ev in pinned:
            f = _ev_fields(ev)
            line = f"{f['ts']} [{f['type']}] {f['host']} {f['user']} — {f['msg']}"
            if f["src"]:
                line += f" (from {f['src']})"
            bullet(line)

    max_flagged = int(tpl.get("max_flagged") or 50)
    if sections.get("flagged", True) and flagged:
        h2(L("flagged_events"))
        for ev in flagged[:max_flagged]:
            f = _ev_fields(ev)
            line = f"{f['ts']} [{f['type']}] {f['host']} {f['user']} — {f['msg']}"
            if f["src"]:
                line += f" (from {f['src']})"
            bullet(line)
        if len(flagged) > max_flagged:
            doc.add_paragraph(L("and_more", n=len(flagged) - max_flagged)).italic = True

    # ── Analyst notes
    notes = data.get("notes") or ""
    if sections.get("notes", True) and notes:
        h2(L("analyst_notes"))
        for ln in notes.splitlines():
            if ln.strip():
                doc.add_paragraph(ln)

    if (tpl.get("footer_md") or "").strip():
        p = doc.add_paragraph(tpl["footer_md"].strip())
        for run in p.runs:
            run.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _now_label() -> str:
    from datetime import UTC, datetime

    return _ts(datetime.now(UTC).isoformat())
